from flask import Flask, request, redirect, url_for, send_from_directory, flash, render_template
from docx import Document
import os

app = Flask(__name__)
app.secret_key = "supersecretkey"
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def add_references_to_word(doc_file_path, references):
    # Charger le document Word
    doc = Document(doc_file_path)
    
    # Rechercher les références existantes dans le document
    existing_references = set(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    
    # Parcourir chaque référence à ajouter
    for reference in references:
        if reference not in existing_references:
            # Ajouter une nouvelle ligne avec la référence
            doc.add_paragraph(reference)
    
    # Sauvegarder le document Word
    doc.save(doc_file_path)

def read_references_from_file(file_path):
    with open(file_path, 'r') as file:
        references = [line.strip() for line in file.readlines() if line.strip()]
    return references

@app.route('/', methods=['GET', 'POST'])
def upload_files():
    if request.method == 'POST':
        word_file = request.files['word_file']
        text_file = request.files['text_file']

        if not word_file or not text_file:
            flash("All fields are required.")
            return redirect(request.url)

        word_path = os.path.join(UPLOAD_FOLDER, "Bibliographie.docx")
        text_path = os.path.join(UPLOAD_FOLDER, "Ref.txt")

        try:
            word_file.save(word_path)
            text_file.save(text_path)
        except Exception as e:
            print(f"Error saving files: {e}")
            flash("There was an error saving the files.")
            return redirect(request.url)

        try:
            references = read_references_from_file(text_path)
            add_references_to_word(word_path, references)
        except Exception as e:
            print(f"Error processing files: {e}")
            flash("There was an error processing the files.")
            return redirect(request.url)

        return redirect(url_for('download_file', filename="Bibliographie.docx"))

    return render_template('upload.html')

@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

if __name__ == '__main__':
    app.run(debug=True)
